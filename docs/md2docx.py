#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
md2docx.py: 将本项目风格的 markdown 转为 .docx
- 标题 (#/##/###) -> Word 标题样式
- 代码块 -> 等宽字体 + 浅灰底纹
- 表格 -> Word 表格
- 列表 -> Word 列表
- ASCII 图 (在代码块内) -> 等宽字体
- 粗体/斜体/行内代码 -> Word run
- 块引用 -> Word 引用样式
"""

import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


def set_cell_shading(cell, fill_hex):
    """设置表格单元格底纹"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def set_cell_borders(cell, color="888888", size="4"):
    """设置表格单元格边框"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tc_pr.append(tcBorders)


def set_run_font(run, name="Arial", size=None, bold=None, italic=None, color=None):
    """设置 run 字体"""
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:cs'), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_code_block_paragraph(doc, code_text):
    """添加一个代码块段落 (整段一个 paragraph, 带灰色底纹)"""
    lines = code_text.rstrip('\n').split('\n')
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.right_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        # 整段底纹
        pPr = p._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F4F4F4')
        pPr.append(shd)
        # 等宽字体
        run = p.add_run(line if line else ' ')
        set_run_font(run, name="Consolas", size=9)


def add_inline_runs(paragraph, text, base_size=11, base_bold=False):
    """在段落里添加内联 run, 处理 **bold** / *italic* / `code`"""
    # 解析顺序: 先 `code`, 再 **bold**, 再 *italic*
    # 简单状态机
    pattern = re.compile(r'(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*|\[[^\]]+\]\([^)]+\))')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            set_run_font(run, name="Microsoft YaHei", size=base_size, bold=base_bold)
        token = m.group(0)
        if token.startswith('`') and token.endswith('`'):
            inner = token[1:-1]
            run = paragraph.add_run(inner)
            set_run_font(run, name="Consolas", size=base_size - 1, color="C7254E")
            # 浅红底
            rPr = run._element.get_or_add_rPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F9F2F4')
            rPr.append(shd)
        elif token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, name="Microsoft YaHei", size=base_size, bold=True)
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Microsoft YaHei", size=base_size, italic=True)
        elif token.startswith('['):
            # 链接 [text](url) - 简化处理,只显示文本
            m2 = re.match(r'\[([^\]]+)\]\([^)]+\)', token)
            if m2:
                run = paragraph.add_run(m2.group(1))
                set_run_font(run, name="Microsoft YaHei", size=base_size, color="0563C1")
                run.font.underline = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, name="Microsoft YaHei", size=base_size, bold=base_bold)


def parse_table_block(lines, start):
    """解析 markdown 表格, 返回 (rows, end_index)"""
    rows = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith('|'):
        line = lines[i].strip()
        # 跳过分隔行 (|---|---|)
        if re.match(r'^\|?[\s\-:|]+\|?$', line) and '---' in line:
            i += 1
            continue
        # 拆 cell
        if line.startswith('|'):
            line = line[1:]
        if line.endswith('|'):
            line = line[:-1]
        cells = [c.strip() for c in line.split('|')]
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc, rows):
    """添加 Word 表格"""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    # 补齐 cell
    rows = [r + [''] * (ncols - len(r)) for r in rows]
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # 清除默认段落
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_inline_runs(p, cell_text, base_size=10, base_bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, "D5E8F0")
            set_cell_borders(cell)


def main():
    src = sys.argv[1] if len(sys.argv) > 1 else 'docs/freertos_tutorial.md'
    dst = sys.argv[2] if len(sys.argv) > 2 else 'docs/freertos_tutorial.docx'

    with open(src, 'r', encoding='utf-8') as f:
        content = f.read()
    lines = content.split('\n')

    doc = Document()
    # 默认字体
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
    rFonts.set(qn('w:hAnsi'), 'Microsoft YaHei')
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    rFonts.set(qn('w:cs'), 'Microsoft YaHei')

    # 标题样式
    for h_level, size in [(1, 18), (2, 15), (3, 13), (4, 12)]:
        try:
            hs = doc.styles[f'Heading {h_level}']
            hs.font.name = 'Microsoft YaHei'
            hs.font.size = Pt(size)
            hs.font.bold = True
            hs.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
            rPr_h = hs.element.get_or_add_rPr()
            rFonts_h = rPr_h.find(qn('w:rFonts'))
            if rFonts_h is None:
                rFonts_h = OxmlElement('w:rFonts')
                rPr_h.insert(0, rFonts_h)
            rFonts_h.set(qn('w:ascii'), 'Microsoft YaHei')
            rFonts_h.set(qn('w:hAnsi'), 'Microsoft YaHei')
            rFonts_h.set(qn('w:eastAsia'), 'Microsoft YaHei')
        except KeyError:
            pass

    # 页边距
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    i = 0
    in_code = False
    code_buf = []
    in_list = False  # 简化: 不维护 list state, 每次遇到空行结束

    while i < len(lines):
        line = lines[i]
        stripped = line.lstrip()

        # ===== 代码块 =====
        if stripped.startswith('```'):
            if in_code:
                # 结束
                add_code_block_paragraph(doc, '\n'.join(code_buf))
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ===== 标题 =====
        if stripped.startswith('####'):
            text = stripped[4:].strip()
            doc.add_heading(text, level=4)
            i += 1
            continue
        if stripped.startswith('###'):
            text = stripped[3:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue
        if stripped.startswith('##'):
            text = stripped[2:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue
        if stripped.startswith('# '):
            text = stripped[2:].strip()
            doc.add_heading(text, level=1)
            i += 1
            continue

        # ===== 水平线 =====
        if stripped == '---':
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '2E75B6')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ===== 表格 =====
        if stripped.startswith('|'):
            rows, i = parse_table_block(lines, i)
            add_table(doc, rows)
            continue

        # ===== 块引用 =====
        if stripped.startswith('> '):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.right_indent = Inches(0.3)
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '24')
            left.set(qn('w:space'), '8')
            left.set(qn('w:color'), '2E75B6')
            pBdr.append(left)
            pPr.append(pBdr)
            add_inline_runs(p, text, base_size=10.5, base_bold=False)
            i += 1
            continue

        # ===== 列表 =====
        if re.match(r'^\s*[-*]\s+', line):
            text = re.sub(r'^\s*[-*]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            add_inline_runs(p, text, base_size=11)
            i += 1
            continue
        if re.match(r'^\s*\d+\.\s+', line):
            text = re.sub(r'^\s*\d+\.\s+', '', line)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(2)
            add_inline_runs(p, text, base_size=11)
            i += 1
            continue

        # ===== 空行 =====
        if not stripped:
            i += 1
            continue

        # ===== 普通段落 =====
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.4
        add_inline_runs(p, line, base_size=11)
        i += 1

    doc.save(dst)
    print(f"已生成: {dst}")


if __name__ == '__main__':
    main()
